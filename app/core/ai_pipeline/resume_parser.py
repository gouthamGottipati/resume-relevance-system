import re
import json
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import spacy
from app.core.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class ContactInfo:
    """Contact information extracted from resume."""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None
    location: Optional[str] = None


@dataclass
class Education:
    """Education information."""
    degree: Optional[str] = None
    institution: Optional[str] = None
    location: Optional[str] = None
    graduation_year: Optional[int] = None
    gpa: Optional[float] = None
    relevant_courses: List[str] = None
    honors: List[str] = None


@dataclass
class WorkExperience:
    """Work experience entry."""
    title: str
    company: str
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    duration_months: Optional[int] = None
    description: List[str] = None
    technologies: List[str] = None
    achievements: List[str] = None


@dataclass
class Project:
    """Project information."""
    title: str
    description: str
    technologies: List[str] = None
    url: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None


@dataclass
class ParsedResume:
    """Complete parsed resume structure."""
    contact_info: ContactInfo
    summary: Optional[str] = None
    skills: List[str] = None
    education: List[Education] = None
    work_experience: List[WorkExperience] = None
    projects: List[Project] = None
    certifications: List[str] = None
    languages: List[str] = None
    awards: List[str] = None
    total_experience_years: Optional[float] = None
    raw_text: str = ""
    parsing_confidence: float = 0.0


class ResumeParser:
    """Advanced resume parser with multiple extraction strategies."""
    
    def __init__(self):
        self.nlp = None
        self._load_nlp_model()
        
        # Regex patterns for extraction
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        self.linkedin_pattern = re.compile(r'linkedin\.com/in/[\w-]+', re.IGNORECASE)
        self.github_pattern = re.compile(r'github\.com/[\w-]+', re.IGNORECASE)
        self.url_pattern = re.compile(r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?')
        
        # Skill categories for better categorization
        self.skill_categories = {
            'programming_languages': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'kotlin', 'swift', 'typescript', 'scala', 'r', 'matlab'],
            'web_technologies': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'bootstrap', 'jquery'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sqlite', 'cassandra', 'dynamodb'],
            'cloud_platforms': ['aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'kubernetes', 'docker'],
            'data_science': ['pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'matplotlib', 'seaborn', 'jupyter'],
            'tools': ['git', 'jenkins', 'ansible', 'terraform', 'nginx', 'apache', 'linux', 'windows', 'macos']
        }
    
    def _load_nlp_model(self):
        """Load spaCy NLP model."""
        try:
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded spaCy English model")
        except OSError:
            logger.warning("spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def parse_pdf(self, file_path: str) -> ParsedResume:
        """Parse PDF resume using multiple strategies."""
        text = ""
        parsing_confidence = 0.0
        
        try:
            # Strategy 1: PyMuPDF (best for most PDFs)
            text_pymupdf = self._extract_text_pymupdf(file_path)
            
            # Strategy 2: pdfplumber (best for tables and complex layouts)
            text_pdfplumber = self._extract_text_pdfplumber(file_path)
            
            # Choose the best extraction
            if len(text_pdfplumber) > len(text_pymupdf):
                text = text_pdfplumber
                parsing_confidence = 0.85
            else:
                text = text_pymupdf
                parsing_confidence = 0.80
                
            logger.info(f"Extracted {len(text)} characters from PDF")
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            parsing_confidence = 0.0
        
        return self._parse_text_content(text, parsing_confidence)
    
    def _extract_text_pymupdf(self, file_path: str) -> str:
        """Extract text using PyMuPDF."""
        text = ""
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return text
    
    def _extract_text_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber."""
        text = ""
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        return text
    
    def parse_docx(self, file_path: str) -> ParsedResume:
        """Parse DOCX resume."""
        text = ""
        parsing_confidence = 0.0
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            parsing_confidence = 0.90  # DOCX parsing is usually reliable
            logger.info(f"Extracted {len(text)} characters from DOCX")
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            parsing_confidence = 0.0
        
        return self._parse_text_content(text, parsing_confidence)
    
    def parse_text(self, text: str) -> ParsedResume:
        """Parse plain text resume."""
        return self._parse_text_content(text, 1.0)
    
    def _parse_text_content(self, text: str, parsing_confidence: float) -> ParsedResume:
        """Parse extracted text content."""
        # Clean and normalize text
        cleaned_text = self._clean_text(text)
        
        # Extract components
        contact_info = self._extract_contact_info(cleaned_text)
        summary = self._extract_summary(cleaned_text)
        skills = self._extract_skills(cleaned_text)
        education = self._extract_education(cleaned_text)
        work_experience = self._extract_work_experience(cleaned_text)
        projects = self._extract_projects(cleaned_text)
        certifications = self._extract_certifications(cleaned_text)
        languages = self._extract_languages(cleaned_text)
        awards = self._extract_awards(cleaned_text)
        
        # Calculate total experience
        total_experience = self._calculate_total_experience(work_experience)
        
        return ParsedResume(
            contact_info=contact_info,
            summary=summary,
            skills=skills,
            education=education,
            work_experience=work_experience,
            projects=projects,
            certifications=certifications,
            languages=languages,
            awards=awards,
            total_experience_years=total_experience,
            raw_text=text,
            parsing_confidence=parsing_confidence
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n', text)
        return text.strip()
    
    def _extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information."""
        contact = ContactInfo()
        
        # Extract email
        email_matches = self.email_pattern.findall(text)
        contact.email = email_matches[0] if email_matches else None
        
        # Extract phone
        phone_matches = self.phone_pattern.findall(text)
        contact.phone = phone_matches[0] if phone_matches else None
        
        # Extract LinkedIn
        linkedin_matches = self.linkedin_pattern.findall(text)
        contact.linkedin = f"https://{linkedin_matches[0]}" if linkedin_matches else None
        
        # Extract GitHub
        github_matches = self.github_pattern.findall(text)
        contact.github = f"https://{github_matches[0]}" if github_matches else None
        
        # Extract name (first few words, excluding common headers)
        lines = text.split('\n')[:5]  # Check first 5 lines
        skip_words = {'resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary'}
        
        for line in lines:
            words = line.strip().split()
            if 2 <= len(words) <= 4:  # Name typically 2-4 words
                line_lower = line.lower()
                if not any(skip in line_lower for skip in skip_words):
                    if not self.email_pattern.search(line) and not self.phone_pattern.search(line):
                        contact.name = line.strip().title()
                        break
        
        return contact
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary/objective."""
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview', 'introduction']
        
        lines = text.split('\n')
        summary_start = None
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in summary_keywords):
                if len(line_lower.split()) <= 3:  # Likely a header
                    summary_start = i + 1
                    break
        
        if summary_start:
            summary_lines = []
            for i in range(summary_start, min(summary_start + 5, len(lines))):
                line = lines[i].strip()
                if line and not self._is_section_header(line):
                    summary_lines.append(line)
                elif summary_lines:  # Stop at next section
                    break
            
            return ' '.join(summary_lines) if summary_lines else None
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical and soft skills."""
        skills = set()
        text_lower = text.lower()
        
        # Extract from skills section
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'technologies', 'competencies'])
        if skills_section:
            # Split by common delimiters
            skill_text = re.split(r'[,;|•\n]', skills_section)
            for skill in skill_text:
                skill = skill.strip().strip('•-*')
                if skill and len(skill.split()) <= 3:  # Skills are usually 1-3 words
                    skills.add(skill.title())
        
        # Extract known skills from entire text
        for category, skill_list in self.skill_categories.items():
            for skill in skill_list:
                if skill.lower() in text_lower:
                    skills.add(skill.title())
        
        return list(skills)
    
    def _extract_education(self, text: str) -> List[Education]:
        """Extract education information."""
        education_list = []
        
        education_section = self._extract_section(text, ['education', 'academic', 'qualifications'])
        if not education_section:
            return education_list
        
        # Split into potential education entries
        entries = re.split(r'\n(?=[A-Z])', education_section)
        
        for entry in entries:
            if len(entry.strip()) > 10:  # Minimum length for valid entry
                edu = self._parse_education_entry(entry)
                if edu.degree or edu.institution:
                    education_list.append(edu)
        
        return education_list
    
    def _parse_education_entry(self, entry: str) -> Education:
        """Parse individual education entry."""
        edu = Education()
        
        # Extract degree
        degree_patterns = [
            r'(bachelor|master|phd|doctorate|associate|diploma|certificate)[\s\w]*',
            r'(b\.?[sa]\.?|m\.?[sa]\.?|ph\.?d\.?|mba)',
            r'(undergraduate|graduate)'
        ]
        
        for pattern in degree_patterns:
            match = re.search(pattern, entry, re.IGNORECASE)
            if match:
                edu.degree = match.group().strip().title()
                break
        
        # Extract institution (typically the longest capitalized phrase)
        words = entry.split()
        potential_institutions = []
        
        current_phrase = []
        for word in words:
            if word and word[0].isupper():
                current_phrase.append(word)
            else:
                if len(current_phrase) >= 2:
                    potential_institutions.append(' '.join(current_phrase))
                current_phrase = []
        
        if current_phrase and len(current_phrase) >= 2:
            potential_institutions.append(' '.join(current_phrase))
        
        if potential_institutions:
            # Choose the longest as most likely to be institution name
            edu.institution = max(potential_institutions, key=len)
        
        # Extract graduation year
        year_matches = re.findall(r'(19|20)\d{2}', entry)
        if year_matches:
            edu.graduation_year = int(year_matches[-1])  # Take the latest year
        
        # Extract GPA
        gpa_match = re.search(r'gpa:?\s*([0-9.]+)', entry, re.IGNORECASE)
        if gpa_match:
            try:
                edu.gpa = float(gpa_match.group(1))
            except ValueError:
                pass
        
        return edu
    
    def _extract_work_experience(self, text: str) -> List[WorkExperience]:
        """Extract work experience."""
        experience_list = []
        
        experience_section = self._extract_section(text, ['experience', 'employment', 'work history', 'professional experience'])
        if not experience_section:
            return experience_list
        
        # Split into potential experience entries
        entries = re.split(r'\n(?=[A-Z].*(?:at|@|\|))', experience_section)
        
        for entry in entries:
            if len(entry.strip()) > 20:  # Minimum length for valid entry
                exp = self._parse_experience_entry(entry)
                if exp.title and exp.company:
                    experience_list.append(exp)
        
        return experience_list
    
    def _parse_experience_entry(self, entry: str) -> WorkExperience:
        """Parse individual work experience entry."""
        lines = [line.strip() for line in entry.split('\n') if line.strip()]
        
        if not lines:
            return WorkExperience(title="", company="")
        
        # First line typically contains title and company
        first_line = lines[0]
        
        # Try to separate title and company
        separators = [' at ', ' @ ', ' - ', ' | ']
        title, company = "", ""
        
        for sep in separators:
            if sep in first_line:
                parts = first_line.split(sep, 1)
                title = parts[0].strip()
                company = parts[1].strip()
                break
        
        if not company:  # Fallback: assume first part is title
            words = first_line.split()
            if len(words) > 1:
                title = ' '.join(words[:len(words)//2])
                company = ' '.join(words[len(words)//2:])
        
        # Extract dates
        date_pattern = r'(\d{1,2}/\d{2,4}|\w+\s+\d{4}|\d{4})'
        dates = []
        
        for line in lines[:3]:  # Check first 3 lines for dates
            date_matches = re.findall(date_pattern, line)
            dates.extend(date_matches)
        
        start_date = dates[0] if len(dates) > 0 else None
        end_date = dates[1] if len(dates) > 1 else "Present"
        
        # Extract description (remaining lines)
        description = []
        for line in lines[1:]:
            if not re.search(date_pattern, line):  # Skip lines with dates
                clean_line = line.strip('•-*')
                if clean_line:
                    description.append(clean_line)
        
        return WorkExperience(
            title=title,
            company=company,
            start_date=start_date,
            end_date=end_date,
            description=description
        )
    
    def _extract_projects(self, text: str) -> List[Project]:
        """Extract project information."""
        projects = []
        
        project_section = self._extract_section(text, ['projects', 'personal projects', 'academic projects'])
        if not project_section:
            return projects
        
        # Split into potential project entries
        entries = re.split(r'\n(?=[A-Z])', project_section)
        
        for entry in entries:
            if len(entry.strip()) > 15:
                project = self._parse_project_entry(entry)
                if project.title:
                    projects.append(project)
        
        return projects
    
    def _parse_project_entry(self, entry: str) -> Project:
        """Parse individual project entry."""
        lines = [line.strip() for line in entry.split('\n') if line.strip()]
        
        if not lines:
            return Project(title="", description="")
        
        title = lines[0].strip('•-*').strip()
        description = ' '.join(lines[1:]) if len(lines) > 1 else ""
        
        # Extract technologies
        tech_keywords = ['technologies:', 'tech stack:', 'built with:', 'using:']
        technologies = []
        
        for line in lines:
            line_lower = line.lower()
            for keyword in tech_keywords:
                if keyword in line_lower:
                    tech_part = line_lower.split(keyword)[1]
                    techs = re.split(r'[,;|]', tech_part)
                    technologies.extend([tech.strip() for tech in techs if tech.strip()])
        
        # Extract URL
        url_match = self.url_pattern.search(entry)
        url = url_match.group() if url_match else None
        
        return Project(
            title=title,
            description=description,
            technologies=technologies,
            url=url
        )
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications."""
        cert_section = self._extract_section(text, ['certifications', 'certificates', 'licenses'])
        if not cert_section:
            return []
        
        # Split by lines and clean
        certs = []
        for line in cert_section.split('\n'):
            clean_cert = line.strip('•-*').strip()
            if clean_cert and len(clean_cert) > 5:
                certs.append(clean_cert)
        
        return certs
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract language skills."""
        lang_section = self._extract_section(text, ['languages', 'language skills'])
        if not lang_section:
            return []
        
        languages = []
        for line in lang_section.split('\n'):
            clean_lang = line.strip('•-*').strip()
            if clean_lang:
                languages.append(clean_lang)
        
        return languages
    
    def _extract_awards(self, text: str) -> List[str]:
        """Extract awards and achievements."""
        awards_section = self._extract_section(text, ['awards', 'achievements', 'honors', 'recognition'])
        if not awards_section:
            return []
        
        awards = []
        for line in awards_section.split('\n'):
            clean_award = line.strip('•-*').strip()
            if clean_award:
                awards.append(clean_award)
        
        return awards
    
    def _extract_section(self, text: str, section_keywords: List[str]) -> Optional[str]:
        """Extract a specific section from text."""
        lines = text.split('\n')
        section_start = None
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in section_keywords):
                if len(line_lower.split()) <= 4:  # Likely a header
                    section_start = i + 1
                    break
        
        if section_start is None:
            return None
        
        # Find section end (next major header or end of text)
        section_lines = []
        for i in range(section_start, len(lines)):
            line = lines[i].strip()
            if self._is_section_header(line) and i > section_start + 2:
                break
            section_lines.append(line)
        
        return '\n'.join(section_lines)
    
    def _is_section_header(self, line: str) -> bool:
        """Check if line is likely a section header."""
        line = line.strip()
        if not line:
            return False
        
        # Common section headers
        headers = [
            'experience', 'education', 'skills', 'projects', 'certifications',
            'awards', 'languages', 'summary', 'objective', 'profile'
        ]
        
        line_lower = line.lower()
        word_count = len(line.split())
        
        return (word_count <= 3 and 
                any(header in line_lower for header in headers))
    
    def _calculate_total_experience(self, work_experience: List[WorkExperience]) -> Optional[float]:
        """Calculate total years of experience."""
        if not work_experience:
            return None
        
        total_months = 0
        current_year = datetime.now().year
        
        for exp in work_experience:
            try:
                # Simple calculation - can be improved
                start_year = self._extract_year_from_date(exp.start_date)
                end_year = self._extract_year_from_date(exp.end_date) or current_year
                
                if start_year and end_year:
                    years_diff = end_year - start_year
                    total_months += max(0, years_diff * 12)
            except:
                continue
        
        return round(total_months / 12, 1) if total_months > 0 else None
    
    def _extract_year_from_date(self, date_str: Optional[str]) -> Optional[int]:
        """Extract year from date string."""
        if not date_str or date_str.lower() == 'present':
            return None
        
        # Extract 4-digit year
        year_match = re.search(r'(19|20)\d{2}', date_str)
        return int(year_match.group()) if year_match else None